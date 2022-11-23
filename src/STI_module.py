# -*- encoding: utf-8 -*-
'''
@File    :   STI_module.py
@Time    :   2022/09/29
@Author  :   ATANG_
@Version :   1.0
@Desc    :   科技情报分析与报告生成（Scientific and Technological Information）
'''


from reading_module import SmartReading
import log_module as log
import wordcloud
from pygtrans import Translate
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx import Document
import imageio
import cn2an
import base64
import datetime
import json
import os
import re
import sys
import warnings
sys.path.append('..')
warnings.filterwarnings('ignore')


CACHE_DIR = '../cache/STI'
# CACHE_DIR = './cache/STI'
WORDCLOUD_BACKGROUND_DIR = '../image/word_cloud.png'
# WORDCLOUD_BACKGROUND_DIR = './image/word_cloud.png'
STOPWORDS_PATH = './stopwords.txt'
# STOPWORDS_PATH = './src/stopwords.txt'

COUNTRY_LIST = ['中国', '英国', '美国', '澳大利亚', '日本', '俄罗斯', '法国']
DOMAIN_LIST = ['数据战略', '5G', '卫星通信', '信息通信']

AN2CN = {'0': '〇', '1': '一', '2': '二', '3': '三', '4': '四',
         '5': '五', '6': '六', '7': '七', '8': '八', '9': '九', }


class STIModule:

    def __init__(self):
        self.trans_client = Translate()
        self.smart_reading = SmartReading()

    def _to_json(self, data):
        return json.dumps(data, ensure_ascii=False, indent=4)

    def _save_cache(self, data, path):
        with open(path, 'w') as f:
            f.write(self._to_json(data))

    # 情报读取
    def read_docs(self, doc_paths: list) -> list:

        def read_doc(path):
            doc = Document(path)
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
            return paragraphs

        docs = []
        doc_names = []
        if isinstance(doc_paths, str):
            docs.append(read_doc(doc_paths))
            doc_names.append(os.path.basename(doc_paths).split('.')[0])
        elif isinstance(doc_paths, list) or isinstance(doc_paths, tuple):
            for path in doc_paths:
                docs.append(read_doc(path))
                doc_names.append(os.path.basename(path).split('.')[0])
        return docs, doc_names

    # 情报翻译
    def translate(self, docs):
        translate_docs = []
        for doc in docs:
            dect = self.trans_client.detect(doc[0])
            if dect.language == 'zh-CN':
                translate_docs.append([p.replace(' ', '') for p in doc])
            else:
                trans = self.trans_client.translate(doc)
                translate_docs.append(
                    [t.translatedText.replace(' ', '') for t in trans])
        return translate_docs

    def _image_to_str(self, image):
        image_byte = base64.b64encode(image)
        image_str = image_byte.decode('ascii')  # byte类型转换为str
        return image_str

    # 词云图生成
    def generate_word_cloud(self, docs):

        def _get_new_word(doc):
            filter_doc = [p for p in doc if len(p) > 100]
            text = '\n'.join(filter_doc)
            res = self.smart_reading.smart_reading(text, 'new_words_finder')
            new_word = json.loads(res)['result']
            return new_word

        mask = imageio.imread(WORDCLOUD_BACKGROUND_DIR)
        stop_file = open(STOPWORDS_PATH)
        excludes = set([l.strip() for l in stop_file.readlines()])
        stop_file.close()
        ls = [_get_new_word(doc) for doc in docs]
        ls = [element for list in ls for element in list]
        txt = " ".join(ls)
        w = wordcloud.WordCloud(width=1000, height=700, background_color='white',
                                font_path="STHeiti Medium.ttc",
                                mask=mask, stopwords=excludes)
        w.generate(txt)
        w.to_file(os.path.join(CACHE_DIR, 'sti_wordcloud.png'))

        image = open(os.path.join(CACHE_DIR, 'sti_wordcloud.png'), 'rb').read()
        return self._image_to_str(image)

    # 单篇情报分析
    def single_analyze(self, doc, doc_name) -> dict:

        # 获取发布时间
        def _get_time():
            dates = []
            for p in doc:
                date = re.search(
                    r"(\d{4}年\d{1,2}月\d{1,2}日)|(\d{4}年\d{1,2}月)|(\d{4}年)", p)
                if date:
                    dates.append(date.group())
            dates.sort(reverse=True)
            for date in dates:
                if int(date[:4]) <= datetime.datetime.today().year:
                    return date
            return ''

        # 获取国家名称
        def _get_country():
            for p in doc:
                for c in COUNTRY_LIST:
                    if c in p:
                        return c
            return ''

        # 获取领域方向
        def _get_domain():
            for p in doc:
                for d in DOMAIN_LIST:
                    if d in p:
                        return d
            return ''

        # 获取关键词
        def _get_keyword():
            filter_doc = [p for p in doc if len(p) > 100]
            text = '\n'.join(filter_doc)
            res = self.smart_reading.smart_reading(text, 'key_extract')
            key_word = json.loads(res)['result'][:5]
            return key_word

        # 获取内容摘要
        def _get_summary():
            filter_doc = [p for p in doc if len(p) > 100]
            text = '\n'.join(filter_doc)
            res = self.smart_reading.smart_reading(text, 'summary')
            summary = json.loads(res)['result']
            return summary

        event = dict()
        event['source'] = doc_name
        event['time'] = _get_time()
        event['country'] = _get_country()
        event['domain'] = _get_domain()
        event['key_word'] = _get_keyword()
        event['summary'] = _get_summary()
        return event

    # 情报分析
    def analyze(self, docs, doc_names) -> dict:
        infos = {'events': {}}
        infos['word_cloud'] = self.generate_word_cloud(docs)
        for doc, doc_name in zip(docs, doc_names):
            event = self.single_analyze(doc, doc_name)
            if event['domain'] not in infos['events']:
                infos['events'][event['domain']] = []
            infos['events'][event['domain']].append(event)
        for domain in infos['events']:
            infos['events'][domain].sort(key=lambda e: e['time'])
        return infos

    # 配置格式
    def style_config(self,
                     document,
                     style_name='default_style',
                     style_id=1,
                     font_size=15,
                     bold=False,
                     italic=False,
                     color=RGBColor(0x00, 0x00, 0x00),
                     alignment_center=False,
                     first_line_indent=0,
                     space_before=0,
                     space_after=0,
                     line_spacing=1,
                     font_name=['Times New Roman', '宋体']):
        style = document.styles.add_style(style_name, style_id)
        style.font.size = Pt(font_size)  # 字体大小
        if bold:    # 字体加粗
            style.font.bold = True
        if italic:  # 字体斜体
            style.font.italic = True
        style.font.color.rgb = color    # 字体颜色
        if alignment_center:    # 字体对齐
            style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        style.paragraph_format.first_line_indent = Pt(
            first_line_indent)    # 首行缩进
        style.paragraph_format.space_before = Pt(space_before)  # 段前间距
        style.paragraph_format.space_after = Pt(space_after)    # 段后间距
        style.paragraph_format.line_spacing = line_spacing  # 行间距
        style.font.name = font_name[0]  # 西文字体
        style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name[1])   # 中文字体
        return style

    # 生成封面
    def generate_cover(self, document):
        def _generate_time():
            t = datetime.datetime.today()
            year = ''.join([AN2CN[c] for c in str(t.year)])
            month = cn2an.an2cn(str(t.month))
            day = cn2an.an2cn(str(t.day))
            return year + '年' + month + '月' + day + '日'

        c_mainbody = self.style_config(document, style_name='c_mainbody',
                                       font_size=8, font_name=['Times New Roman', '宋体'])
        document.add_paragraph('\n'*2, style=c_mainbody)
        c_style1 = self.style_config(document, style_name='c_style1',
                                     font_size=18, first_line_indent=36,
                                     line_spacing=1.5, font_name=['Times New Roman', '楷体_GB2312'])
        document.add_paragraph('国防科技战略先导计划', style=c_style1)
        document.add_paragraph('\n'*8, style=c_mainbody)
        c_style2 = self.style_config(document, style_name='c_style2',
                                     font_size=70, color=RGBColor(0xff, 0x00, 0x00),
                                     alignment_center=True, line_spacing=1,
                                     font_name=['Times New Roman', '方正苏轼行书 简繁'])
        document.add_paragraph('网络信息参考', style=c_style2)
        c_style3 = self.style_config(document, style_name='c_style3',
                                     font_size=20, alignment_center=True, space_before=20,
                                     line_spacing=1.5, font_name=['Times New Roman', '楷体_GB2312'])
        document.add_paragraph('（2022年第4季度汇编）', style=c_style3)
        document.add_paragraph('\n'*25, style=c_mainbody)
        c_style4 = self.style_config(document, style_name='c_style4',
                                     font_size=18,  line_spacing=1.5, alignment_center=True,
                                     font_name=['Times New Roman', '楷体_GB2312'])
        document.add_paragraph('军事科学院军事科学信息研究中心', style=c_style4)
        date = _generate_time()
        document.add_paragraph(date, style=c_style4)
        document.add_page_break()
        document.add_page_break()

    # 生成目录
    def generate_directory(self, document, infos):
        d_style1 = self.style_config(document, style_name='d_style1',
                                     font_size=28, alignment_center=True,
                                     space_before=7, space_after=7,
                                     font_name=['Times New Roman', '黑体'])
        document.add_paragraph('目录', style=d_style1)
        d_style2 = self.style_config(document, style_name='d_style2',
                                     font_size=16, alignment_center=True,
                                     space_before=16, space_after=16, line_spacing=1.8,
                                     font_name=['Times New Roman', '方正小标宋简体'])
        d_style3 = self.style_config(document, style_name='d_style3',
                                     font_size=12, line_spacing=1.8,
                                     font_name=['Times New Roman', '宋体'])
        for id, domain in enumerate(infos['events']):
            document.add_paragraph('第{}篇  {}'.format(
                AN2CN[str(id+1)], domain), style=d_style2)
            for event in infos['events'][domain]:
                document.add_paragraph('·{}发布《{}》'.format(
                    event['country'], event['source']), style=d_style3)
        document.add_page_break()

    # 生成正文
    def generate_content(self, document, infos):
        t_style1 = self.style_config(document, style_name='t_style1',
                                     font_size=28, font_name=['Times New Roman', '黑体'])
        t_style2 = self.style_config(document, style_name='t_style2',
                                     font_size=22, alignment_center=True,
                                     line_spacing=1.5, font_name=['Times New Roman', '方正小标宋简体'])
        t_style3 = self.style_config(document, style_name='t_style3',
                                     font_size=24, alignment_center=True,
                                     line_spacing=1.5, font_name=['Times New Roman', '方正小标宋简体'])
        t_style4 = self.style_config(document, style_name='t_style4',
                                     font_size=18, alignment_center=True,
                                     line_spacing=1.7, font_name=['Times New Roman', '黑体'])
        t_style5 = self.style_config(document, style_name='t_style5',
                                     font_size=16, first_line_indent=32,
                                     line_spacing=1.7, font_name=['Times New Roman', '楷体_GB2312'])
        t_style6 = self.style_config(document, style_name='t_style6',
                                     font_size=16, line_spacing=1.7,
                                     font_name=['Times New Roman', '黑体'])
        for id, domain in enumerate(infos['events']):
            document.add_paragraph('\n'*8, style=t_style1)
            document.add_paragraph('第{}篇'.format(
                AN2CN[str(id+1)]), style=t_style2)
            document.add_paragraph(domain, style=t_style3)
            document.add_page_break()
            for event in infos['events'][domain]:
                document.add_paragraph('{}发布《{}》'.format(
                    event['country'], event['source']), style=t_style4)
                document.add_paragraph('{}（{}）'.format(
                    event['summary'], event['time']), style=t_style5)
                document.add_paragraph('关键词：{}'.format(
                    ' '.join(event['key_word'])), style=t_style6)
                document.add_paragraph('', style=t_style4)
            document.add_page_break()

    # 插入词云图
    def insert_wordcloud_picture(self, document):
        p_style1 = self.style_config(document, style_name='p_style1',
                                     font_size=28, alignment_center=True,
                                     space_before=7, space_after=7,
                                     font_name=['Times New Roman', '黑体'])
        document.add_paragraph('语言分布', style=p_style1)
        document.add_paragraph('\n'*2, style=p_style1)
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run("")
        run.add_picture(os.path.join(
            CACHE_DIR, 'sti_wordcloud.png'), width=Cm(15))
        document.add_page_break()

    # 报告生成
    def generate_report(self, infos):
        document = Document()
        self.generate_cover(document)
        self.generate_directory(document, infos)
        self.insert_wordcloud_picture(document)
        self.generate_content(document, infos)
        document.save(os.path.join(CACHE_DIR, '网络信息参考.docx'))

    def STI_manager(self, doc_paths):
        warnings.filterwarnings('ignore')

        docs, doc_names = self.read_docs(doc_paths)
        trans_docs = self.translate(docs)
        self._save_cache(trans_docs, os.path.join(
            CACHE_DIR, 'trans_docs.json'))

        infos = self.analyze(trans_docs, doc_names)
        self._save_cache(infos, os.path.join(CACHE_DIR, 'infos.json'))
        log.INFO('情报分析完成')

        self.generate_report(infos)
        log.INFO('报告生成完成')


if __name__ == '__main__':
    pass
